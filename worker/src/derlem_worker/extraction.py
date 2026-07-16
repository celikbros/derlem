"""Belge ekstraksiyonu: PDF/DOCX -> satır başına belge TXT.

Yüklenen ikili belgeler kanonik metin formatına burada çevrilir. Kural:
paragraflar normalize edilir (iç boşluk/yeni satır tek boşluğa iner) ve
paragraf sınırlarında bölünerek en fazla ``chunk_chars`` karakterlik
belgeler hâlinde, satır başına bir belge olacak şekilde LF ile yazılır.
Görüntü tabanlı (metinsiz) PDF'ler v1'de desteklenmez; OCR ayrı iştir.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Iterator
from zipfile import BadZipFile, LargeZipFile, ZipFile, is_zipfile

EXTRACTION_VERSION = "text-extraction-v1"
SUPPORTED_SUFFIXES = {".pdf", ".docx"}
DEFAULT_CHUNK_CHARS = 4000
DEFAULT_MAX_SOURCE_BYTES = 100 * 1024 * 1024
DEFAULT_MAX_DOCX_ENTRIES = 2048
DEFAULT_MAX_DOCX_UNCOMPRESSED_BYTES = 256 * 1024 * 1024
DEFAULT_MAX_PDF_PAGES = 1000
DEFAULT_MAX_OUTPUT_CHARS = 32 * 1024 * 1024
MEDIA_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

_WHITESPACE_RE = re.compile(r"\s+")


class ExtractionError(ValueError):
    """Belgeden kullanılabilir metin çıkarılamadı."""


@dataclass(frozen=True)
class ExtractionLimits:
    """Parser'a girmeden ve ayrıştırma sırasında uygulanan kaynak sınırları."""

    max_source_bytes: int = DEFAULT_MAX_SOURCE_BYTES
    max_docx_entries: int = DEFAULT_MAX_DOCX_ENTRIES
    max_docx_uncompressed_bytes: int = DEFAULT_MAX_DOCX_UNCOMPRESSED_BYTES
    max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES
    max_output_chars: int = DEFAULT_MAX_OUTPUT_CHARS

    def __post_init__(self) -> None:
        for name, value in (
            ("max_source_bytes", self.max_source_bytes),
            ("max_docx_entries", self.max_docx_entries),
            ("max_docx_uncompressed_bytes", self.max_docx_uncompressed_bytes),
            ("max_pdf_pages", self.max_pdf_pages),
            ("max_output_chars", self.max_output_chars),
        ):
            if value <= 0:
                raise ValueError(f"{name} must be positive")


@dataclass(frozen=True)
class ExtractionReport:
    method: str
    suffix: str
    paragraph_count: int
    document_count: int
    total_chars: int


def needs_extraction(filename: str | None) -> bool:
    if not filename:
        return False
    return Path(filename).suffix.lower() in SUPPORTED_SUFFIXES


def media_type_for(filename: str) -> str:
    return MEDIA_TYPES.get(Path(filename).suffix.lower(), "application/octet-stream")


def _validate_source(path: Path, suffix: str, limits: ExtractionLimits) -> None:
    try:
        source_size = path.stat().st_size
    except OSError as error:
        raise ExtractionError("Belge kaynağı okunamıyor.") from error
    if source_size > limits.max_source_bytes:
        raise ExtractionError(
            "Belge extraction kaynak boyutu sınırını aşıyor "
            f"({source_size} > {limits.max_source_bytes} byte)."
        )

    try:
        with path.open("rb") as source:
            magic = source.read(8)
    except OSError as error:
        raise ExtractionError("Belge kaynağı okunamıyor.") from error

    if suffix == ".pdf":
        if not magic.startswith(b"%PDF-"):
            raise ExtractionError("Dosya uzantısı PDF ancak PDF imzası bulunamadı.")
        return
    if suffix == ".docx":
        if not magic.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
            raise ExtractionError("Dosya uzantısı DOCX ancak ZIP imzası bulunamadı.")
        _validate_docx_container(path, limits)
        return
    raise ExtractionError(f"Desteklenmeyen belge türü: {suffix}")


def _validate_docx_container(path: Path, limits: ExtractionLimits) -> None:
    try:
        if not is_zipfile(path):
            raise ExtractionError("DOCX ZIP kapsayıcısı bozuk veya eksik.")
        with ZipFile(path) as archive:
            entries = archive.infolist()
            if len(entries) > limits.max_docx_entries:
                raise ExtractionError(
                    "DOCX ZIP girdi sayısı sınırını aşıyor "
                    f"({len(entries)} > {limits.max_docx_entries})."
                )
            total_uncompressed = 0
            names: set[str] = set()
            for entry in entries:
                total_uncompressed += entry.file_size
                if total_uncompressed > limits.max_docx_uncompressed_bytes:
                    raise ExtractionError(
                        "DOCX açılmış içerik boyutu sınırını aşıyor "
                        f"({total_uncompressed} > "
                        f"{limits.max_docx_uncompressed_bytes} byte)."
                    )
                if entry.flag_bits & 0x1:
                    raise ExtractionError("Şifreli DOCX ZIP girdileri desteklenmiyor.")
                if entry.filename in names:
                    raise ExtractionError(
                        f"DOCX ZIP tekrarlı girdi içeriyor: {entry.filename}"
                    )
                names.add(entry.filename)
            required = {"[Content_Types].xml", "word/document.xml"}
            if not required.issubset(names):
                raise ExtractionError("ZIP kapsayıcısı geçerli bir DOCX belgesi değil.")
            corrupt_entry = archive.testzip()
            if corrupt_entry is not None:
                raise ExtractionError(f"DOCX ZIP girdisi bozuk: {corrupt_entry}")
    except ExtractionError:
        raise
    except (BadZipFile, LargeZipFile, OSError, RuntimeError) as error:
        raise ExtractionError("DOCX ZIP kapsayıcısı bozuk veya okunamıyor.") from error


def _pdf_paragraphs(path: Path, *, max_pages: int) -> Iterator[str]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ExtractionError("Şifreli PDF desteklenmiyor; şifresini kaldırıp yeniden yükleyin.")
    page_count = len(reader.pages)
    if page_count > max_pages:
        raise ExtractionError(
            f"PDF sayfa sınırını aşıyor ({page_count} > {max_pages})."
        )
    for page in reader.pages:
        text = page.extract_text() or ""
        for block in re.split(r"\n\s*\n", text):
            yield block


def _docx_paragraphs(path: Path) -> Iterator[str]:
    from docx import Document

    document = Document(str(path))
    for paragraph in document.paragraphs:
        yield paragraph.text
    for table in document.tables:
        for row in table.rows:
            yield " | ".join(cell.text for cell in row.cells)


def extract_paragraphs(
    path: Path,
    suffix: str,
    *,
    limits: ExtractionLimits | None = None,
) -> Iterator[str]:
    suffix = suffix.lower()
    effective_limits = limits or ExtractionLimits()
    _validate_source(path, suffix, effective_limits)
    if suffix == ".pdf":
        source = _pdf_paragraphs(path, max_pages=effective_limits.max_pdf_pages)
    elif suffix == ".docx":
        source = _docx_paragraphs(path)
    else:
        raise ExtractionError(f"Desteklenmeyen belge türü: {suffix}")
    for raw in source:
        normalized = _WHITESPACE_RE.sub(" ", raw).strip()
        if normalized:
            yield normalized


def _split_oversized_paragraph(paragraph: str, chunk_chars: int) -> Iterator[str]:
    """Split one normalized paragraph without exceeding ``chunk_chars``.

    Paragraph boundaries remain the preferred chunk boundary, but a single
    oversized paragraph must not violate the extraction contract. Prefer the
    last whitespace inside the limit and fall back to a hard character split
    for exceptionally long tokens.
    """
    start = 0
    paragraph_length = len(paragraph)
    while paragraph_length - start > chunk_chars:
        boundary = start + chunk_chars
        split_at = paragraph.rfind(" ", start, boundary + 1)
        if split_at <= start:
            split_at = boundary
        segment = paragraph[start:split_at].strip()
        if segment:
            yield segment
        start = split_at
        while start < paragraph_length and paragraph[start].isspace():
            start += 1
    if start < paragraph_length:
        yield paragraph[start:].strip()


def convert_file(
    source_path: Path,
    target_path: Path,
    *,
    suffix: str,
    chunk_chars: int = DEFAULT_CHUNK_CHARS,
    limits: ExtractionLimits | None = None,
) -> ExtractionReport:
    """Belgeyi satır-başına-belge TXT'ye çevirir ve raporunu döndürür."""
    if chunk_chars <= 0:
        raise ValueError("chunk_chars must be positive")
    paragraph_count = 0
    document_count = 0
    total_chars = 0
    chunk: list[str] = []
    chunk_size = 0
    extracted_chars = 0
    effective_limits = limits or ExtractionLimits()

    try:
        with target_path.open("wb") as target:

            def flush() -> None:
                nonlocal chunk, chunk_size, document_count, total_chars
                if not chunk:
                    return
                line = " ".join(chunk)
                target.write(line.encode("utf-8") + b"\n")
                document_count += 1
                total_chars += len(line)
                chunk = []
                chunk_size = 0

            for paragraph in extract_paragraphs(
                source_path, suffix, limits=effective_limits
            ):
                paragraph_count += 1
                extracted_chars += len(paragraph)
                if extracted_chars > effective_limits.max_output_chars:
                    raise ExtractionError(
                        "Belgeden çıkarılan toplam metin sınırı aşılıyor "
                        f"({extracted_chars} > {effective_limits.max_output_chars} karakter)."
                    )
                for segment in _split_oversized_paragraph(paragraph, chunk_chars):
                    addition = len(segment) if not chunk else len(segment) + 1
                    if chunk and chunk_size + addition > chunk_chars:
                        flush()
                        addition = len(segment)
                    chunk.append(segment)
                    chunk_size += addition
            flush()
    except ExtractionError:
        target_path.unlink(missing_ok=True)
        raise
    except Exception as error:
        target_path.unlink(missing_ok=True)
        raise ExtractionError(
            "Belge ayrıştırılamadı; dosya bozuk veya desteklenmeyen yapıdadır."
        ) from error

    if document_count == 0:
        target_path.unlink(missing_ok=True)
        raise ExtractionError(
            "Belgeden metin çıkarılamadı; dosya görüntü tabanlı (taranmış) olabilir. "
            "OCR bu sürümde desteklenmiyor."
        )
    return ExtractionReport(
        method=EXTRACTION_VERSION,
        suffix=suffix.lower(),
        paragraph_count=paragraph_count,
        document_count=document_count,
        total_chars=total_chars,
    )
